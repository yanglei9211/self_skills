#!/usr/bin/env python3
"""Convert our news-summary markdown to a .docx file.

Handles:
  - # / ## / ### headings
  - Bullet lists with inline **bold**, `code`, and [text](url) hyperlinks
  - Blockquotes (>)
  - Horizontal rules (---)
  - Chinese-friendly default font (SimSun + 微软雅黑 fallback on Windows,
    PingFang on macOS)
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_hyperlink(paragraph, text, url):
    """Insert a clickable hyperlink run into `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F6FEB")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Inline parser: walks a line and emits runs of (kind, payload)
# kinds: 'text', 'bold', 'code', 'link'
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"              # **bold**
    r"|(`[^`]+`)"                   # `code`
    r"|(\[[^\]]+\]\([^)]+\))"       # [text](url)
)


def parse_inline(text):
    tokens = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()]))
        g = m.group(0)
        if g.startswith("**"):
            tokens.append(("bold", g[2:-2]))
        elif g.startswith("`"):
            tokens.append(("code", g[1:-1]))
        else:
            tm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", g)
            tokens.append(("link", (tm.group(1), tm.group(2))))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    return tokens


def render_inline(para, tokens):
    for kind, payload in tokens:
        if kind == "text":
            para.add_run(payload)
        elif kind == "bold":
            r = para.add_run(payload)
            r.bold = True
        elif kind == "code":
            r = para.add_run(payload)
            r.font.name = "Menlo"
            r.font.size = Pt(10)
        elif kind == "link":
            text, url = payload
            add_hyperlink(para, text, url)


def set_chinese_font(run, size=11):
    """Apply a CJK-friendly font to a run (works on macOS + Windows)."""
    run.font.size = Pt(size)
    run.font.name = "Calibri"  # latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    # East Asia font preference
    rFonts.set(qn("w:eastAsia"), "PingFang SC")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:ascii"), "Calibri")


def convert(md_path, docx_path):
    doc = Document()

    # Set a default CJK-friendly style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "PingFang SC")

    with open(md_path) as f:
        lines = f.read().split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.rstrip()

        if not s.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}\s*$", s):
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        if s.startswith("# "):
            h = doc.add_heading(s[2:].strip(), level=1)
            i += 1
            continue
        if s.startswith("## "):
            h = doc.add_heading(s[3:].strip(), level=2)
            i += 1
            continue
        if s.startswith("### "):
            h = doc.add_heading(s[4:].strip(), level=3)
            i += 1
            continue

        # Blockquote
        if s.lstrip().startswith("> "):
            content = s.lstrip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run("▎ ")
            r.font.color.rgb = RGBColor(0x6A, 0x73, 0x7D)
            render_inline(p, parse_inline(content))
            i += 1
            continue

        # Bullet list
        if s.lstrip().startswith("- "):
            content = s.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, parse_inline(content))
            i += 1
            continue

        # Numbered list (not really used in our file but just in case)
        if re.match(r"^\s*\d+\.\s", s):
            content = re.sub(r"^\s*\d+\.\s", "", s)
            p = doc.add_paragraph(style="List Number")
            render_inline(p, parse_inline(content))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        render_inline(p, parse_inline(s))
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
